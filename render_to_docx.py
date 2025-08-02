from docx import Document
from typing import Any, Union

def render_summary_to_docx_old(summary: dict, output_path: str = "voynich_summary.docx"):
    """
    Renders the given summary dictionary into a .docx file.
    
    Args:
        summary (dict): The summary as returned by summarize_all_responses().
        output_path (str): Path where the .docx will be saved.
    """
    doc = Document()
    doc.add_heading("Voynich Manuscript Summary", level=1)
    
    # 1. Summary Hypothesis
    doc.add_heading("1. Summary Hypothesis", level=2)
    doc.add_paragraph(summary.get("summary_hypothesis", ""))

    # 2. Key Patterns
    doc.add_heading("2. Key Patterns", level=2)
    for pattern in summary.get("key_patterns", []):
        doc.add_paragraph(pattern, style="List Bullet")

    # 3. Possible Language Influences
    doc.add_heading("3. Possible Language Influences", level=2)
    for lang in summary.get("possible_language_influences", []):
        doc.add_paragraph(lang, style="List Bullet")

    # 4. Conclusion Type
    doc.add_heading("4. Conclusion Type", level=2)
    doc.add_paragraph(summary.get("conclusion_type", ""))

    # 5. Confidence Range
    doc.add_heading("5. Confidence Range", level=2)
    cr = summary.get("confidence_range", [])
    if isinstance(cr, list) and len(cr) == 2:
        doc.add_paragraph(f"{cr[0]} to {cr[1]}")
    else:
        doc.add_paragraph("")

    # 6. Open Questions
    doc.add_heading("6. Open Questions", level=2)
    for question in summary.get("open_questions", []):
        doc.add_paragraph(question, style="List Number")

    # Save the document
    doc.save(output_path)
    print(f"Document saved to {output_path}")


def render_dict_to_docx(data: dict, output_path: str = "voynich_generic_summary.docx"):
    """
    Renders any nested dictionary into a structured .docx file using its keys as headings.
    
    Args:
        data (dict): The dictionary to render.
        output_path (str): Path where the .docx will be saved.
    """
    doc = Document()
    doc.add_heading("Generated Summary", level=1)

    def render_section(content: Union[dict, list, str, int, float, None], level: int = 2):
        """
        Recursively renders content based on its type.
        """
        if isinstance(content, dict):
            for key, value in content.items():
                doc.add_heading(str(key), level=level)
                render_section(value, level=min(level + 1, 5))  # limit heading levels
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, (str, int, float)):
                    doc.add_paragraph(str(item), style="List Bullet")
                elif isinstance(item, dict):
                    render_section(item, level=level)
                else:
                    doc.add_paragraph(str(item))
        elif content is None:
            doc.add_paragraph("(none)")
        else:
            doc.add_paragraph(str(content))

    render_section(data)
    doc.save(output_path)
    print(f"Document saved to {output_path}")